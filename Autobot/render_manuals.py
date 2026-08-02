# -*- coding: utf-8 -*-
"""Regera os .docx e .pdf dos manuais a partir dos .md.

Os .md sao a fonte da verdade. Se so eles forem editados, o comprador abre um
.pdf desatualizado -- pior do que nao ter editado nada.

Tipografia por idioma: chines, japones e coreano precisam de fontes CID que o
reportlab embute (STSong-Light, HeiseiKakuGo-W5, HYSMyeongJo-Medium). Cirilico
e turco saem em Arial do proprio Windows, que cobre os dois alfabetos. Sem
isso o PDF vira um bloco de quadradinhos.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (KeepTogether, ListFlowable, ListItem,
                                PageBreak, Paragraph, SimpleDocTemplate,
                                Spacer, Table, TableStyle)
from reportlab.lib import colors

ROOT = Path(r"C:\Users\Lucas Pedroso\OneDrive\SantoGral\Santo Gral"
            r"\venda mql5\Pacote_Completo_White_Rabbit_X\Languages")
WINFONTS = Path(r"C:\Windows\Fonts")

# Fonte CID embutida no reportlab para cada escrita logografica.
CID_FONTS = {
    "03_Chinese": "STSong-Light",
    "05_Japanese": "HeiseiKakuGo-W5",
    "07_Korean": "HYSMyeongJo-Medium",
}
BRAND = colors.HexColor("#1E3A5F")
MUTED = colors.HexColor("#5A6B82")


def register_latin() -> tuple[str, str]:
    """Arial cobre latino estendido, cirilico e turco."""
    try:
        pdfmetrics.registerFont(TTFont("WR", str(WINFONTS / "arial.ttf")))
        pdfmetrics.registerFont(TTFont("WR-Bold", str(WINFONTS / "arialbd.ttf")))
        pdfmetrics.registerFontFamily("WR", normal="WR", bold="WR-Bold")
        return "WR", "WR-Bold"
    except Exception as exc:  # noqa: BLE001 - queremos o fallback, nao o traceback
        print(f"  Arial indisponivel ({exc}); usando Helvetica")
        return "Helvetica", "Helvetica-Bold"


def fonts_for(lang: str, latin: tuple[str, str]) -> tuple[str, str]:
    cid = CID_FONTS.get(lang)
    if not cid:
        return latin
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(cid))
        return cid, cid
    except Exception as exc:  # noqa: BLE001
        print(f"  Fonte CID {cid} indisponivel ({exc}); usando latina")
        return latin


def inline(md: str) -> str:
    """Marcacao inline do Markdown -> tags do reportlab."""
    md = (md.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    md = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", md)
    md = re.sub(r"`(.+?)`", r'<font face="Courier">\1</font>', md)
    md = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", md)
    return md


def plain(md: str) -> str:
    md = re.sub(r"\*\*(.+?)\*\*", r"\1", md)
    md = re.sub(r"`(.+?)`", r"\1", md)
    md = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", md)
    return md


def parse(markdown: str) -> list[tuple[str, object]]:
    """Markdown -> lista de blocos (tipo, conteudo)."""
    blocks: list[tuple[str, object]] = []
    buffer: list[str] = []
    bullets: list[str] = []
    table: list[list[str]] = []

    def flush_text() -> None:
        if buffer:
            blocks.append(("p", " ".join(buffer).strip()))
            buffer.clear()

    def flush_bullets() -> None:
        if bullets:
            blocks.append(("ul", list(bullets)))
            bullets.clear()

    def flush_table() -> None:
        if table:
            blocks.append(("table", [list(r) for r in table]))
            table.clear()

    for raw in markdown.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_text(); flush_bullets(); flush_table()
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_text(); flush_bullets()
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                continue  # linha separadora
            table.append(cells)
            continue
        flush_table()
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush_text(); flush_bullets()
            blocks.append((f"h{len(m.group(1))}", m.group(2)))
            continue
        if stripped in ("---", "***", "___"):
            flush_text(); flush_bullets()
            blocks.append(("hr", ""))
            continue
        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            flush_text()
            bullets.append(m.group(1))
            continue
        m = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if m:
            flush_text()
            bullets.append(m.group(2))
            continue
        if stripped.startswith(">"):
            flush_text(); flush_bullets()
            blocks.append(("quote", stripped.lstrip("> ").strip()))
            continue
        buffer.append(stripped)

    flush_text(); flush_bullets(); flush_table()
    return blocks


def to_docx(blocks: list[tuple[str, object]], target: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10.5)

    for kind, content in blocks:
        if kind.startswith("h") and kind != "hr":
            level = int(kind[1])
            p = doc.add_heading(plain(str(content)), level=min(level, 4))
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        elif kind == "p":
            doc.add_paragraph(plain(str(content)))
        elif kind == "ul":
            for item in content:  # type: ignore[union-attr]
                doc.add_paragraph(plain(item), style="List Bullet")
        elif kind == "quote":
            p = doc.add_paragraph(plain(str(content)))
            p.paragraph_format.left_indent = Pt(18)
            for run in p.runs:
                run.italic = True
        elif kind == "table":
            rows = content  # type: ignore[assignment]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    if c < len(table.columns):
                        table.cell(r, c).text = plain(cell)
        elif kind == "hr":
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(target)


def to_pdf(blocks: list[tuple[str, object]], target: Path,
           regular: str, bold: str) -> None:
    base = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=base["BodyText"], fontName=regular,
                          fontSize=9.5, leading=14, alignment=TA_LEFT,
                          spaceAfter=6)
    quote = ParagraphStyle("quote", parent=body, leftIndent=14,
                           textColor=MUTED, borderPadding=4)
    heads = {
        1: ParagraphStyle("h1", parent=base["Heading1"], fontName=bold,
                          fontSize=17, leading=21, textColor=BRAND,
                          spaceBefore=4, spaceAfter=10),
        2: ParagraphStyle("h2", parent=base["Heading2"], fontName=bold,
                          fontSize=13, leading=17, textColor=BRAND,
                          spaceBefore=12, spaceAfter=6),
        3: ParagraphStyle("h3", parent=base["Heading3"], fontName=bold,
                          fontSize=11, leading=15, textColor=BRAND,
                          spaceBefore=9, spaceAfter=4),
    }
    heads[4] = heads[3]
    heads[5] = heads[3]
    heads[6] = heads[3]

    story: list = []
    for kind, content in blocks:
        if kind.startswith("h") and kind != "hr":
            level = min(int(kind[1]), 6)
            story.append(Paragraph(inline(str(content)), heads[level]))
        elif kind == "p":
            story.append(Paragraph(inline(str(content)), body))
        elif kind == "quote":
            story.append(Paragraph(inline(str(content)), quote))
        elif kind == "ul":
            items = [ListItem(Paragraph(inline(i), body), leftIndent=12)
                     for i in content]  # type: ignore[union-attr]
            story.append(ListFlowable(items, bulletType="bullet",
                                      start="•", leftIndent=14))
        elif kind == "table":
            rows = content  # type: ignore[assignment]
            if not rows:
                continue
            cell = ParagraphStyle("cell", parent=body, fontSize=8.5, leading=11,
                                  spaceAfter=0)
            head = ParagraphStyle("cellh", parent=cell, fontName=bold,
                                  textColor=colors.white)
            data = [[Paragraph(inline(c), head if r == 0 else cell)
                     for c in row] for r, row in enumerate(rows)]
            width = (A4[0] - 40 * mm) / max(len(rows[0]), 1)
            table = Table(data, colWidths=[width] * len(rows[0]), hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), BRAND),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C4D4")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#F2F5F9")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(KeepTogether(table))
            story.append(Spacer(1, 8))
        elif kind == "hr":
            story.append(Spacer(1, 10))

    SimpleDocTemplate(
        str(target), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=target.stem, author="Lucas Pedroso",
    ).build(story)


def main() -> int:
    latin = register_latin()
    langs = sorted(d.name for d in ROOT.iterdir()
                   if d.is_dir() and re.match(r"^\d\d_", d.name))
    made_docx = made_pdf = failed = 0

    for lang in langs:
        regular, bold = fonts_for(lang, latin)
        print(f"{lang} (fonte: {regular})")
        for md in sorted((ROOT / lang).glob("*.md")):
            if md.name.lower() == "readme.md":
                continue
            blocks = parse(md.read_text(encoding="utf-8-sig"))
            try:
                to_docx(blocks, md.with_suffix(".docx"))
                made_docx += 1
            except Exception as exc:  # noqa: BLE001
                print(f"  docx falhou em {md.name}: {exc}")
                failed += 1
            try:
                to_pdf(blocks, md.with_suffix(".pdf"), regular, bold)
                made_pdf += 1
            except Exception as exc:  # noqa: BLE001
                print(f"  pdf falhou em {md.name}: {exc}")
                failed += 1

    print(f"\n.docx gerados: {made_docx}")
    print(f".pdf gerados:  {made_pdf}")
    if failed:
        print(f"falhas: {failed}")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
